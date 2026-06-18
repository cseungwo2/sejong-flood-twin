# -*- coding: utf-8 -*-
"""docx 메모에 3차 결정(건물용도색·교량·슬라이더15m·물표면) 근거 기입."""
import os
from docx import Document

DOC = r"C:\Users\user\Desktop\새 폴더 (4)\1차 데이터_세종시 제공\세종_하천_기본수심_근거.docx"
doc = Document(DOC)

def para(t, b=False):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = b; return p
def bullets(items):
    for t in items: doc.add_paragraph(t, style="List Bullet")
def table(headers, rows):
    tb = doc.add_table(rows=1, cols=len(headers)); tb.style = "Light Grid Accent 1"
    for i, x in enumerate(headers):
        c = tb.rows[0].cells[i]; c.text = ""; run = c.paragraphs[0].add_run(x); run.bold = True
    for row in rows:
        cs = tb.add_row().cells
        for i, v in enumerate(row): cs[i].text = str(v)
    return tb

doc.add_heading("H. 건물 용도색·교량·슬라이더·물표면 (2026-06-18 3차)", level=1)

para("H-1. 건물 용도별 색 (토글, 도시계획 표준색)", True)
bullets([
    "원본 F_FAC_BUILDING(세종)의 USABILITY 필드 = 건축법 시행령 별표1 건물 용도 5자리 코드. 세종 건물 77,425동 중 33%(25,758동)에 용도 입력됨.",
    "geojson과 원본 SHP의 피처 순서가 달라, 용도를 정확히 입히려 SHP에서 buildings.geojson을 재생성(형상+높이+use). 앱이 지형에 건물을 안착하므로 지반고는 불필요.",
    "코드 분포: 01000 단독주택(14,546) 최다, 04000 제2종근생(2,261), 21000 동식물관련(2,040), 18000 창고(1,836), 17000 공장(1,747), 03000 제1종근생(1,551), 02000 공동주택(556).",
])
para("도시계획(용도지역) 표준색 관례에 맞춘 카테고리 색:", True)
table(["카테고리", "포함 용도", "색"],
      [["주거", "단독·공동주택", "노랑"],
       ["상업·근생", "1·2종근생, 판매, 업무, 숙박, 위락", "빨강"],
       ["공업", "공장, 위험물, 발전", "보라"],
       ["창고·물류", "창고, 자동차, 자원순환", "연보라"],
       ["농업·동식물", "동물·식물관련", "초록"],
       ["공공·교육", "교육, 의료, 문화, 종교, 공공, 운동 등", "파랑"],
       ["용도 미상", "미입력(67%)", "회색"]])
bullets([
    "기본 OFF(평소 흰색). 토글 ON 시 용도색 — '어느 지역 주민이 침수 피해 큰지' 파악용. 어지러움 방지로 끌 수 있게.",
    "건물 호버 시 용도 카테고리명 + 침수 여부 + 지반높이 표시(코드 번호는 숨김).",
])

para("H-2. 교량·고가도로 (OSM + 지형 기반 높이)", True)
bullets([
    "OSM(오버패스 API)에서 세종 교량 2,322개(이름 915개: 금강자전거길·세종로·감성3교·경부고속선·이응다리류) 추출, 길이≥20m 또는 이름 있는 1,757개 사용.",
    "⚠️ OSM엔 교량 상판 높이(height) 태그가 0개 — 정확한 높이 데이터 없음.",
    "높이 해법: '다리는 양 끝에서 도로면에 연결되고 골짜기로 내려가지 않는다'는 원리 → 양 끝점의 DEM 지형고를 직선 보간해 상판 높이로 사용(+1m 노면). 별도 높이자료 없이 위경도+지형으로 도로에 이어지고 강/계곡 위로 부상하는 모델.",
    "한계: 고가도로(도로 위 도로)는 끝점 지형이 램프 하부라 다소 낮게 추정될 수 있음. 정밀화하려면 교량 관리대장(상판 표고) 필요.",
])

para("H-3. 침수심 슬라이더·물표면", True)
bullets([
    "범위 0~15m로 확대(극한 대비), step 0.05m.",
    "수위를 목표치로 매 프레임 부드럽게 수렴(ease)시켜 '갑자기 띵 나타나는' 이질감 제거.",
    "물 표면 얕은 가장자리를 smoothstep으로 페이드해 경계 어색함 완화.",
    "라벨 토글 끄기 버그 수정(렌더 루프가 모두 꺼지면 라벨 컨테이너 숨김).",
])

para("H-4. UI·운영", True)
bullets([
    "화면 설명문구 제거(부제·무방비 가정·tick 등 — 정직한 근거는 본 문서에 보관, 화면은 깔끔하게).",
    "왼쪽 패널: 스크롤 + 섹션 접기/펴기, 버튼 많은 '홍수위험지도' 섹션 기본 접힘.",
    "start.bat: python 탐지를 'py -3' 우선으로 보강(실행 안 되던 문제 대응).",
])

doc.save(DOC)
print("APPENDED session3 ->", DOC, "| paras:", len(doc.paragraphs), "tables:", len(doc.tables))
