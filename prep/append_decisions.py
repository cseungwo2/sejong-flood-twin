# -*- coding: utf-8 -*-
"""기존 근거 문서(docx)에 '트윈 구축 결정·근거 로그'를 추가(append). 메모장처럼 누적."""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_BREAK

DOC = r"C:\Users\user\Desktop\새 폴더 (4)\1차 데이터_세종시 제공\세종_하천_기본수심_근거.docx"
doc = Document(DOC)

def h(t, lvl=1):
    p = doc.add_heading(t, level=lvl)
    return p

def para(t, bold=False):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = bold; return p

def bullets(items):
    for t in items:
        doc.add_paragraph(t, style="List Bullet")

def table(headers, rows):
    tb = doc.add_table(rows=1, cols=len(headers)); tb.style = "Light Grid Accent 1"
    for i, x in enumerate(headers):
        c = tb.rows[0].cells[i]; c.text = ""; run = c.paragraphs[0].add_run(x); run.bold = True
    for row in rows:
        cs = tb.add_row().cells
        for i, v in enumerate(row): cs[i].text = str(v)
    return tb

# 페이지 구분 후 부록 시작
doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
t = doc.add_heading("부록 — 디지털트윈 구축 결정·근거 로그", level=0)
for r in t.runs: r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
para("이 부록은 트윈을 만들며 내린 결정과 그 근거를 누적 기록하는 메모다. (업데이트: 2026-06-18)")

# 1. 하천망 레이어
h("A. 하천망 레이어 (하천구역·소하천구역)", 1)
bullets([
    "원본 SHP: 하천유역_세종(LSMD_CONT_UJ201, Polygon 60, EPSG:5174), 소하천_세종(LSMD_CONT_UJ301, Polygon 237, EPSG:5186).",
    "MNUM 코드로 분류: UJB1=하천구역(실제 하천 45) / UJB4=홍수관리구역(규제토지 15, 제외) / UJC1=소하천구역(실제 141) / UJC2=소하천예정지(미통수 96, 제외).",
    "물 표현은 실제 하천(UJB1+UJC1)만. 좌표 5174/5186 → 4326 변환 시 datum transform 적용(Korean_1985_To_WGS_1984_1 / KGD2002_To_WGS_1984_1) — 안 하면 지형과 수십m 어긋남.",
    "국가하천은 미호천 1개뿐(금강 본류는 이 데이터에 없음). '금강천'은 동명 지류라 지방하천(부분매칭 오분류 주의).",
    "ALIAS(하천명)는 CP949 인코딩 — arcpy가 깨뜨려 dbf 직접 디코드 필요.",
])
para("색상(서로·다른 레이어와 안 겹치게):", True)
table(["등급", "색(hex)"], [["국가하천", "0x1f86d0 (진청록 파랑)"],
                            ["지방하천", "0x49b0e6 (중간 파랑)"],
                            ["소하천", "0x8fe3ea (밝은 아쿠아)"]])
para("시점 회전 시 수면이 지형과 깊이충돌(깜빡임) → polygonOffset(Factor/Units −4~−6)로 깊이 우선권 고정해 해결.")

# 2. 세종 경계 클리핑
h("B. 세종시 행정경계 클리핑", 1)
bullets([
    "트윈이 직사각형 bbox로 보이던 것을 세종시 실제 형상으로 잘라냄.",
    "경계원본: 연속지적도 LSMD_CONT_LDREG_세종(206,880 필지, EPSG:5186). 필지 전체를 래스터화하면 도로·하천 지목까지 포함돼 '채워진 세종 형상'이 됨(단일 경계선 파일 불필요).",
    "마스크 sejong_mask.png(1281×1472, 흰=내부) 생성 → 지형 셰이더 onBeforeCompile로 경계 밖 픽셀 discard. 위성영상은 지형 텍스처라 지형+영상 동시에 잘림.",
])

# 3. 침수도(홍수위험지도) 경량화·색
h("C. 침수·범람 오버레이 경량화 및 색상", 1)
bullets([
    "원본 하천범람 geojson이 43~70MB(합 214MB)로 과중 → dissolve + SimplifyPolygon('1.5 Meters') + 좌표 6자리(~0.1m)로 각 0.6~0.8MB(합 2.8MB)로 축소. 영역·형상은 유지.",
    "500년 도시침수(500_SJ.shp): 16,987 '침수 필지'(JIBUN/PNU 단위, 침수심 등급 없음) → dissolve로 침수영역 병합(16MB→1.4MB).",
    "도시침수 30/50/80/100년: 세종 범위 필지 단위, 동일 방식으로 경량화(각 0.9~1.3MB).",
])
para("색 체계(계열로 구분, 하천망 파랑과 안 겹치게):", True)
table(["레이어군", "색 계열"], [
    ["하천범람 100/200/500/기왕최대", "따뜻한 램프(노랑0xffd24a→주황→빨강0xff5e5e→보라0xc060ff)"],
    ["도시침수 30/50/80/100", "초록 램프(연두0xa7e86b→0x5fcf52→0x2aa247→진초록0x10783a)"],
    ["도시침수 500", "보라 0x9b5de5 (GIS 보라영역과 일치)"],
])

# 4. HAND 침수심 모델 (핵심)
h("D. 침수심 시뮬레이션 — HAND 모델 (핵심 결정)", 1)
para("문제: 기존 침수심바는 '도시 전체 단일 수평면'을 올리는 방식이라, 가장 낮은 국가하천(금강·미호)만 먼저 잠기고 지방하천·소하천은 반영되지 않았다.")
para("해결: HAND(Height Above Nearest Drainage) 도입.", True)
bullets([
    "정의: 각 지점이 '흘러드는 가장 가까운 하천'보다 수직으로 몇 m 높은지.",
    "계산(arcpy/Spatial Analyst): DEM(10m, 5186) → Fill → FlowDirection → 하천망(하천구역+소하천구역) 래스터화 → FlowDistance(distance_type=VERTICAL) = HAND.",
    "terrain.bin과 동일하게 4배 블록평균 축소 → hand.bin(634×906, terrain과 셀 일치). 배수 안되는 셀(전체의 4.4%)은 99999 → 영원히 안 잠김.",
    "슬라이더 의미 변경: '하천 수면이 평소보다 차오른 높이(m)' (0~15m). 침수심 = 슬라이더 − HAND. 수면고도 = 지형 + max(침수심,0).",
    "효과: 국가·지방·소하천이 각자 하상에서 동시에 차오르고, 옆으로 번지는 거리는 그 지점 DEM 경사로 결정(평지 넓게, 산 좁게).",
    "건물: 건물 위치의 HAND를 샘플링해 건물별 침수 수면을 따로 적용(벽면 침수심 음영). 통계(침수건물 수·면적·최대침수심)도 HAND 기준으로 재계산.",
])
para("한계(정직하게 — 과장 금지):", True)
bullets([
    "정적 모델이다. 진짜 유체 흐름·시간전파·운동량은 미반영.",
    "건물이 댐처럼 물을 막는 효과, 제방 붕괴 타이밍 미반영 — 이는 HEC-RAS 2D(다음 단계) 영역.",
    "40m 격자 해상도 한계로 좁은 소하천 하상은 근사치.",
    "즉 'DEM 지형 기반 정적 침수'로는 타당하나 '완전한 물리(유체역학) 시뮬레이션'은 아니다.",
])

# 5. 기타 운영 메모
h("E. 운영 메모", 1)
bullets([
    "셰이프파일은 .shp/.shx/.dbf/.prj 4종 세트가 있어야 arcpy가 연다(.shp만으론 안 됨).",
    "arcpy는 git-bash의 /tmp 경로를 못 읽음 → 윈도우 경로 사용.",
    "라이브 편집 시 캐시: index.html의 app.js?v=N 버전을 올려 무력화. 무캐시 서버(nocache_server.py) 단일 인스턴스만 사용(이중서버 금지).",
    "공유: cloudflared 퀵터널(trycloudflare.com)로 임시 링크 — PC/터널 켜진 동안만 유효, 주소 매번 바뀜. 영구는 Netlify에 dist/ 동기화.",
])

doc.save(DOC)
print("APPENDED ->", DOC, "| paras:", len(doc.paragraphs), "tables:", len(doc.tables))
