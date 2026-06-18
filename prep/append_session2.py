# -*- coding: utf-8 -*-
"""docx 메모에 색상·UI·슬라이더·침수해석 결정(2차) 기입."""
import os
from docx import Document

DOC = r"C:\Users\user\Desktop\새 폴더 (4)\1차 데이터_세종시 제공\세종_하천_기본수심_근거.docx"
doc = Document(DOC)

def para(t, b=False):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = b; return p
def bullets(items):
    for t in items: doc.add_paragraph(t, style="List Bullet")
def table(headers, rows):
    tb = doc.add_table(rows=1, cols=len(headers)); tb.style = "Light Grid Accent 1"
    for i, x in enumerate(headers):
        c = tb.rows[0].cells[i]; c.text = ""; run = c.paragraphs[0].add_run(x); run.bold = True
    for row in rows:
        cs = tb.add_row().cells
        for i, v in enumerate(row): cs[i].text = str(v)
    return tb

doc.add_heading("G. 색상·슬라이더·침수 해석·UI 결정 (2026-06-18 2차)", level=1)

para("G-1. 침수심 슬라이더 범위 — 국내 극한 도시침수 기준", True)
bullets([
    "슬라이더 = '하천 수위가 평소보다 차오른 높이(m)'. 범위 0~7m, 0.05m 단위.",
    "근거: 국내 역대 극한 도시침수 = 2023.7.15 청주 오송 궁평2지하차도 침수(미호강 범람·제방 붕괴, 약 6만 톤 유입, 지하차도 완전 침수 ~6m, 14명 사망). 이 미호강이 곧 세종의 국가하천.",
    "지표면(거리·주거지) 극한은 대략 1.5~2.5m, 저지대·지하공간은 ~6m 수준.",
    "→ 극한 6m + 여유 1m = 최대 7m로 설정. step 0.05m로 잘게(조금만 올려도 1m로 점프하던 문제 해소).",
    "출처: 궁평2지하차도 침수 사고(위키백과/나무위키), 국무조정실 보도자료(2023.7.28).",
])

para("G-2. '1m에 수천 가구 침수, 맞나?' — HAND로 검증", True)
para("우리 HAND 데이터로 직접 계산한 결과(전체 건물 76,981동 중):")
table(["수위 상승", "침수 건물", "비율", "침수면적"],
      [["+0.5m", "4,133", "5%", "49.5 km²"],
       ["+1m", "5,426", "7%", "55.7 km²"],
       ["+2m", "8,088", "11%", "66.6 km²"],
       ["+3m", "10,565", "14%", "76.6 km²"]])
bullets([
    "건물 HAND 중앙값 12m(절반은 하천보다 12m 이상 높아 안전). 하천보다 1m 이내로 낮은 건물 = 7%.",
    "해석: 버그로 전체가 잠기는 게 아니라, 잠기는 5,400여 동은 실제로 하천변 저지대(범람원)에 있는 낮은 건물들 → 지형상 타당.",
    "단, 제방·배수펌프·우수관거 같은 방어시설과 40m DEM이 못 잡는 좁은 제방을 미반영한 '무방비 최대가정'이라 실제보다 과장됨.",
    "→ (A) 패널에 '무방비 가정·제방/배수 미반영' 명시. (C) 추후 10m 해상도 HAND로 제방 반영 예정.",
])

para("G-3. 색상 체계 변경", True)
bullets([
    "'기왕최대 하천범람'(river_max) 오버레이 제거.",
    "도시침수 색을 초록 램프 → 선명한 크림슨/적색 계열 '서로 구분되는' 색으로(눈에 띄게, 단색 그라데이션 지양).",
])
table(["도시침수", "색(hex)"],
      [["30년", "0xff5c8a 핑크"], ["50년", "0xe5383b 레드"], ["80년", "0xb5179e 마젠타"],
       ["100년", "0x9d0208 다크레드"], ["500년", "0x6a040f 와인"]])
bullets([
    "하천범람(100/200/500)은 따뜻한 노랑~빨강 유지. 하천망(국가/지방/소하천)은 파랑/청록 유지.",
    "침수 오버레이 불투명도 0.42→0.6(더 또렷). 시점 회전 시 찌글거림(z-fighting)은 polygonOffset(−4)로 제거(오버레이·물 메시 모두).",
])

para("G-4. 건물·UI 정리", True)
bullets([
    "건물: 데이터 지반고(gz) 대신 렌더 지형 높이에 안착(땅에 박힘 해소) + 양면 렌더(지붕 안 덮히는 문제 해소).",
    "건물 호버 문구: 'EL/미침수' → '안 잠김 · 지반높이 32m' / '침수 ≈ 1.8m 잠김' 쉬운 말로.",
    "왼쪽 패널: 길어서 아래 잘리던 것 → 스크롤 + 섹션 헤더 클릭으로 접기/펴기.",
    "중앙 하단 침수심 패널: 폭·글자·여백 축소(컴팩트).",
])

para("G-5. 향후(데이터 필요)", True)
bullets([
    "C: HAND를 10m 해상도로 재계산해 제방·둑 반영(침수 과장 완화).",
    "#4: 건물 용도별(주거/상업 등) 레이어 — 건물 데이터 용도 필드 확인 필요.",
    "#5: 고가도로·교량(이응다리 등, 지면+~12m) — OSM(오픈스트리트맵) 교량 데이터로 구현 가능.",
])

doc.save(DOC)
print("APPENDED session2 ->", DOC, "| paras:", len(doc.paragraphs), "tables:", len(doc.tables))
