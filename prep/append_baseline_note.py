# -*- coding: utf-8 -*-
"""docx 메모에 '기본수심 vs HAND 넘침 기준' 근거 섹션을 추가(append)."""
import os
from docx import Document

DOC = r"C:\Users\user\Desktop\새 폴더 (4)\1차 데이터_세종시 제공\세종_하천_기본수심_근거.docx"
doc = Document(DOC)

def para(t, bold=False):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = bold; return p

def bullets(items):
    for t in items:
        doc.add_paragraph(t, style="List Bullet")

def table(headers, rows):
    tb = doc.add_table(rows=1, cols=len(headers)); tb.style = "Light Grid Accent 1"
    for i, x in enumerate(headers):
        c = tb.rows[0].cells[i]; c.text = ""; run = c.paragraphs[0].add_run(x); run.bold = True
    for row in rows:
        cs = tb.add_row().cells
        for i, v in enumerate(row): cs[i].text = str(v)
    return tb

doc.add_heading("F. 기본수심(0.7/1.5/2.5m) vs HAND 넘침 기준 — 무엇을 침수 판정에 썼나", level=1)
para("두 값은 목적이 다르며, 침수심 시뮬레이션의 '넘침(범람) 판정'에는 기본수심이 아니라 HAND(실제 둑 높이)를 사용했다.")
table(["구분", "기본수심 0.7/1.5/2.5m", "HAND 넘침 기준(현재 채택)"],
      [["의미", "평상시 채널에 담긴 물 깊이(표시용)", "그 지점의 실제 둑 높이(제방고) — 하천보다 땅이 얼마나 높은지"],
       ["출처", "하천 등급별 대표값(경험칙)", "DEM 지형에서 지점별 실측"],
       ["용도", "평상시 하천에 물 보이게 하는 시각 기준", "물이 둑을 넘는 순간 = 범람 판정"]])

para("왜 HAND(실제 둑 높이)가 더 정확한가 — 근거:", True)
bullets([
    "범람은 본질적으로 '수위가 그 자리 둑(제방·지반)을 넘느냐'의 문제다. 둑 높이는 구간마다 다르다.",
    "0.7/1.5/2.5m는 등급별 평균값이라, 둑이 낮은 구간은 과소(실제보다 늦게 범람), 둑이 높은 구간은 과대(너무 빨리 범람)로 어긋난다.",
    "HAND는 각 지점의 실제 DEM 둑 높이를 쓰므로 '어디서 먼저 넘치는가'가 지형대로 정확하다. 평지는 조금만 차도 넓게, 산비탈은 많이 차도 안 넘침이 자동 반영된다.",
])
para("현재 구현 상태:", True)
bullets([
    "침수 판정 = 슬라이더 상승량 > 그 지점 HAND(실제 둑 높이) → 그 차이만큼 침수. 국가/지방/소하천 모두 각자 하상에서 동시 적용.",
    "기본수심 0.7/1.5/2.5m는 아직 모델 입력이 아니라, 평상시 채널 물 표시용 기준으로만 보관 중.",
])
para("향후 옵션(원하면 연결):", True)
bullets([
    "슬라이더 0 = 하천이 이미 기본수심(0.7/1.5/2.5m)만큼 찬 상태로 두고, 그 위로 상승해 (기본수심+상승)이 둑을 넘으면 범람으로 연결 가능. 이때 기본수심이 '평상시 수위'로 모델에 들어간다.",
    "단, 넘침 임계는 그 경우에도 실제 둑(HAND) 기준이 가장 정확하다.",
])

doc.save(DOC)
print("APPENDED baseline-vs-HAND note ->", DOC, "| paras:", len(doc.paragraphs), "tables:", len(doc.tables))
