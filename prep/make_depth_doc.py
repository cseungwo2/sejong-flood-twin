# -*- coding: utf-8 -*-
"""세종 디지털트윈 - 하천 등급별 기본 수심 근거 문서(docx) 생성."""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "docs")
os.makedirs(OUT_DIR, exist_ok=True)
OUT = os.path.join(OUT_DIR, "세종_하천_기본수심_근거.docx")

doc = Document()

# 기본 글꼴(한글)
style = doc.styles["Normal"]
style.font.name = "맑은 고딕"
style.font.size = Pt(10.5)

def h(text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def para(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    return p

def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, htxt in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        run = c.paragraphs[0].add_run(htxt)
        run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return t

# ===== 표지 =====
title = doc.add_heading("세종시 침수 디지털트윈\n하천 등급별 기본 수심(Baseline Water Depth) 설정 근거", level=0)
for r in title.runs:
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
sub = doc.add_paragraph("REDILab.UBFlooding  ·  작성일 2026-06-18")
sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
doc.add_paragraph()

# ===== 1. 목적 =====
h("1. 목적", 1)
para("침수 디지털트윈에서 평상시(침수 슬라이더 0 상태)에도 하천에 자연스러운 물을 표현하기 위해, "
     "하천 등급별 '기본 수심(baseline water depth)'을 합리적 근거를 갖고 설정한다. "
     "이 값은 ① 평상시 하천 수면 표시의 기준이자 ② 침수 시뮬레이션의 시작 수위(0이 아닌 baseline)로 사용된다.")

# ===== 2. 대상 데이터 =====
h("2. 대상 데이터(SHP)", 1)
para("바탕화면 보유 SHP 2종을 분석. 둘 다 폴리곤이며 좌표계가 달라 디지털트윈(EPSG:4326)으로 변환이 필요하다.")
table(
    ["레이어(코드)", "파일", "도형", "전체", "좌표계"],
    [
        ["하천유역(UJ201)", "LSMD_CONT_UJ201_5174_36_202606.shp", "Polygon", "60", "EPSG:5174"],
        ["소하천(UJ301)", "LSMD_CONT_UJ301_36_202606.shp", "Polygon", "237", "EPSG:5186"],
    ],
)
doc.add_paragraph()
para("속성(MNUM) 코드로 세분류한 결과 — 실제 물이 흐르는 면과 규제·계획 구역이 섞여 있어 분리가 필요하다.", bold=True)
table(
    ["코드", "정체", "개수", "기본 수심 적용"],
    [
        ["UJB1", "하천구역 (실제 하천)", "45", "O (국가/지방 수심)"],
        ["UJB4", "홍수관리구역 (규제 토지)", "15", "X (상시 수면 아님)"],
        ["UJC1", "소하천구역 (실제 소하천)", "141", "O (소하천 수심)"],
        ["UJC2", "소하천예정지 (미개수 계획)", "96", "X (아직 미통수)"],
    ],
)
doc.add_paragraph()
para("※ ALIAS(하천명) 필드는 EUC-KR(CP949) 인코딩으로, 복원 시 미호천·조천·행화천·제천 등 실제 하천명이 확인됨.")

# ===== 3. 하천 등급 분류 =====
h("3. 세종시 하천 등급 분류", 1)
para("하천법상 등급(국가/지방하천)과 소하천정비법상 소하천으로 구분된다. 세종 관내 분류는 다음과 같다.")
table(
    ["등급", "관리", "세종 관내 대상"],
    [
        ["국가하천", "환경부장관", "금강(본류), 미호천(미호강)"],
        ["지방하천", "시·도지사", "조천, 월하천, 행화천, 제천 등 (UJB1 중 비국가)"],
        ["소하천", "시장·군수", "소하천구역 141개소 (UJC1)"],
    ],
)
para("※ 보유 SHP 속성에는 국가/지방 구분 필드가 없어, 하천명 기준으로 미호천(+금강)만 국가하천, 나머지 하천구역은 지방하천으로 분류한다.")

# ===== 4. 기본 수심 도출 =====
h("4. 등급별 기본 수심 도출", 1)
para("도출 전제(중요): '하천 등급별 평균 수심'이라는 단일 공표 통계는 존재하지 않는다. "
     "평수위는 표고(EL.m)값이며 수심=평수위−하상고로 구간마다 다르고, 실측치는 하천기본계획 단면자료(RIMGIS)에 분산되어 있다. "
     "따라서 ① 평수위 개념(연중 185일 이상 유지되는 보통수위), ② 등급별 하폭·유량 규모 차이, "
     "③ 실무 경험칙 범위(국가·지방 1.5~3m, 소하천 0.5~1.2m)를 종합해 '대표 기본값'을 설정하고, 추후 실측으로 교체 가능하게 둔다.")
doc.add_paragraph()
table(
    ["등급", "적용 대상", "기본 수심(m)", "근거 / 비고"],
    [
        ["국가하천", "미호천 (+금강)", "2.5", "대하천(미호강 L=89.2km, 유역 1,860.9㎢). 경험칙 1.5~3m의 상단."],
        ["지방하천", "조천·월하천 등 (UJB1 비국가)", "1.5", "중소 지방하천. 경험칙 범위 하단(보수적)."],
        ["소하천", "소하천구역 (UJC1)", "0.7", "경험칙 0.5~1.2m의 중앙값 근방(보수적)."],
        ["(제외)", "홍수관리구역(UJB4)·소하천예정지(UJC2)", "—", "상시 수면이 아니므로 기본 수심 미적용."],
    ],
)
doc.add_paragraph()
para("이 값들은 디지털트윈 코드에 '등급별 파라미터'로 분리 저장되어, 실측치 확보 시 즉시 교체·조정할 수 있다.")

# ===== 5. 한계 및 향후 정밀화 =====
h("5. 한계 및 향후 정밀화", 1)
for t in [
    "정적 bathtub 표현으로, 흐름·시간전파는 미반영(동적은 HEC-RAS 2D 연동이 다음 단계).",
    "기본 수심은 등급 대표값이며 구간별 실제 수심과 차이가 있을 수 있음.",
    "정밀화 경로: RIMGIS/하천기본계획의 종·횡단 단면에서 구간별 평수위(EL.m)·하상고를 확보하면 reach 단위 실측 수위로 교체.",
    "40m DEM 해상도 한계로 좁은 소하천 하상은 근사치.",
]:
    doc.add_paragraph(t, style="List Bullet")

# ===== 6. 출처 =====
h("6. 출처", 1)
sources = [
    ("평수위 정의(연 185일 보통수위) — 하천용어해설, 국토교통부 부산지방국토관리청",
     "https://www.molit.go.kr/brocm/USR/WPGE0201/m_24277/DTL.jsp"),
    ("금강수계 하천기본계획 보고서(금강·미호천·갑천·유등천) — 건설기술정보시스템(CODIL)",
     "https://www.codil.or.kr/viewDtlConRpt.do?gubun=rpt&pMetaCode=OTKCRK200688"),
    ("하천정보관리시스템(RIMGIS) — 하천기본계획·계획홍수위 등 성과품",
     "https://www.river.go.kr/"),
    ("계획홍수위 조회 — 하천정보관리시스템",
     "https://www.river.go.kr/basicPlan/designFloodLevel.do"),
    ("미호강(구 미호천) 제원(L=89.2km, 유역 1,860.9㎢) — 한국민족문화대백과사전",
     "https://encykorea.aks.ac.kr/Article/E0020049"),
    ("미호강 개요 — 나무위키",
     "https://namu.wiki/w/미호강"),
    ("소하천정비법 — 국가법령정보센터",
     "https://www.law.go.kr/LSW/lsInfoP.do?lsiSeq=61918"),
    ("소하천정비종합계획 사례(장성군) — 공공데이터포털",
     "https://www.data.go.kr/data/15067889/fileData.do"),
    ("하천기본계획 수립지침 2023 — 환경부",
     "https://www.law.go.kr/LSW/flDownload.do?flSeq=134143703"),
    ("WAMIS 국가수자원관리종합정보시스템",
     "https://www.wamis.go.kr/"),
]
for i, (label, url) in enumerate(sources, 1):
    p = doc.add_paragraph(style="List Number")
    p.add_run(label + "\n")
    r = p.add_run(url)
    r.font.color.rgb = RGBColor(0x0B, 0x57, 0xD0)
    r.font.size = Pt(9)

doc.save(OUT)
print("SAVED:", OUT)
